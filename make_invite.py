# make_invite.py - Generates a Word document with custom invitations from a
# .txt list.
import docx


#text1 = str(input("Insert first line text for invitation: "))
doc = docx.Document('invite.docx')
count = 1
with open('lista.txt', 'r') as invite:
    for line in invite.readlines():
        doc.add_paragraph("I would be a pleasure to have you")
        doc.paragraphs[count].style = 'Invite'
        count += 1

        doc.add_paragraph(line)
        doc.paragraphs[count].style = 'Invite Name'
        count += 1

        doc.add_paragraph("at 11010 Memory Lane on the evening of")
        doc.paragraphs[count].style = 'Invite'
        count += 1

        doc.add_paragraph('April 1st')
        doc.paragraphs[count].style = 'Invite Date'
        count += 1

        doc.add_paragraph('at 7 o\'clock')
        doc.paragraphs[count].style = 'Invite Date'

        doc.paragraphs[count].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
        count += 1

    doc.save('new_invite.docx')
